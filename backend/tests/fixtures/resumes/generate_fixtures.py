"""Regenerates the binary fixture files in this directory.

These are synthetically constructed resumes, not real ones — no real
resume corpus is available yet (Phase 10 brings in real, larger corpora
per docs/DATASET_STRATEGY.md). They exist to exercise layout diversity
(single-column PDF, two-column PDF, DOCX, corrupt files of each type) per
the Phase 1 test procedure in docs/ROADMAP.md, with known strings and
known layout signals (font size, bold, heading styles) the tests assert
against.

Run manually after changing this file:
    python tests/fixtures/resumes/generate_fixtures.py
"""
from pathlib import Path

import docx
import pymupdf

OUT_DIR = Path(__file__).parent


def build_single_column_pdf(path: Path) -> None:
    pdf = pymupdf.open()
    page = pdf.new_page(width=612, height=792)
    y = 60
    page.insert_text((50, y), "Jordan Ellis", fontsize=20, fontname="hebo")
    y += 22
    page.insert_text((50, y), "jordan.ellis@example.com | (555) 123-4567 | Seattle, WA", fontsize=10)
    y += 30
    page.insert_text((50, y), "TECHNICAL SKILLS", fontsize=13, fontname="hebo")
    y += 16
    page.insert_text((50, y), "Python, SQL, Docker, FastAPI, PyTorch", fontsize=10)
    y += 30
    page.insert_text((50, y), "EXPERIENCE", fontsize=13, fontname="hebo")
    y += 16
    page.insert_text((50, y), "Software Engineer, Acme Corp -- Jun 2021 to Present", fontsize=10)
    y += 14
    page.insert_text((50, y), "- Built a resume parsing pipeline processing 10k documents per day", fontsize=10)
    y += 14
    page.insert_text((50, y), "- Reduced API latency by 40 percent through query optimization", fontsize=10)
    y += 30
    page.insert_text((50, y), "EDUCATION", fontsize=13, fontname="hebo")
    y += 16
    page.insert_text((50, y), "B.S. Computer Science, University of Washington, 2021", fontsize=10)
    pdf.save(path)
    pdf.close()


def build_two_column_pdf(path: Path) -> None:
    pdf = pymupdf.open()
    page = pdf.new_page(width=612, height=792)
    page.insert_text((50, 60), "Taylor Morgan", fontsize=20, fontname="hebo")

    y = 100
    page.insert_text((50, y), "CONTACT", fontsize=12, fontname="hebo")
    y += 16
    page.insert_text((50, y), "taylor.morgan@example.com", fontsize=10)
    y += 14
    page.insert_text((50, y), "Portland, OR", fontsize=10)
    y += 30
    page.insert_text((50, y), "SKILLS", fontsize=12, fontname="hebo")
    y += 16
    page.insert_text((50, y), "Java, Spring Boot, Kubernetes", fontsize=10)
    y += 14
    page.insert_text((50, y), "PostgreSQL, AWS", fontsize=10)

    y = 100
    page.insert_text((320, y), "EXPERIENCE", fontsize=12, fontname="hebo")
    y += 16
    page.insert_text((320, y), "Backend Engineer, Globex Inc -- Mar 2020 to Present", fontsize=10)
    y += 14
    page.insert_text((320, y), "- Migrated monolith to microservices on Kubernetes", fontsize=10)
    y += 14
    page.insert_text((320, y), "- Led a team of 4 engineers", fontsize=10)
    y += 30
    page.insert_text((320, y), "EDUCATION", fontsize=12, fontname="hebo")
    y += 16
    page.insert_text((320, y), "M.S. Software Engineering, Portland State University", fontsize=10)
    pdf.save(path)
    pdf.close()


def build_varied_headers_pdf(path: Path) -> None:
    """Header wording that varies from the other fixtures' ("Technical
    Skills" vs "Core Competencies", "Experience" vs "Employment History"),
    plus one header ("Why Hire Me") that isn't in the section detector's
    alias table at all, to exercise its layout-only fallback path.
    """
    pdf = pymupdf.open()
    page = pdf.new_page(width=612, height=792)
    y = 60
    page.insert_text((50, y), "Morgan Lee", fontsize=20, fontname="hebo")
    y += 22
    page.insert_text((50, y), "morgan.lee@example.com | Denver, CO", fontsize=10)
    y += 30
    page.insert_text((50, y), "PROFESSIONAL SUMMARY", fontsize=13, fontname="hebo")
    y += 16
    page.insert_text((50, y), "Backend engineer with 5 years building distributed systems.", fontsize=10)
    y += 30
    page.insert_text((50, y), "CORE COMPETENCIES", fontsize=13, fontname="hebo")
    y += 16
    page.insert_text((50, y), "Go, Kafka, Terraform, GCP", fontsize=10)
    y += 30
    page.insert_text((50, y), "WHY HIRE ME", fontsize=13, fontname="hebo")
    y += 16
    page.insert_text((50, y), "A relentless focus on reliability and mentoring junior engineers.", fontsize=10)
    y += 30
    page.insert_text((50, y), "EMPLOYMENT HISTORY", fontsize=13, fontname="hebo")
    y += 16
    page.insert_text((50, y), "Staff Engineer, Initrode -- 2019 to Present", fontsize=10)
    y += 30
    page.insert_text((50, y), "ACADEMIC BACKGROUND", fontsize=13, fontname="hebo")
    y += 16
    page.insert_text((50, y), "B.S. Computer Science, Colorado State University", fontsize=10)
    pdf.save(path)
    pdf.close()


def build_resume_docx(path: Path) -> None:
    document = docx.Document()
    document.add_heading("Sam Rivera", level=1)
    document.add_paragraph("sam.rivera@example.com | (555) 987-6543 | Austin, TX")
    document.add_heading("Skills", level=2)
    document.add_paragraph("Go, gRPC, Terraform, GCP, Redis")
    document.add_heading("Experience", level=2)
    document.add_paragraph("Site Reliability Engineer, Initech -- Jan 2019 to Present")
    document.add_paragraph(
        "Cut incident response time by 60 percent by building an on-call runbook system",
        style="List Bullet",
    )
    document.add_heading("Education", level=2)
    document.add_paragraph("B.S. Computer Engineering, University of Texas at Austin, 2018")
    document.save(path)


def build_corrupt_files(pdf_path: Path, docx_path: Path) -> None:
    pdf_path.write_bytes(b"%PDF-1.4 not actually a valid pdf body \x00\x01\x02\xff")
    docx_path.write_bytes(b"this is not a zip file, definitely not OOXML")


if __name__ == "__main__":
    build_single_column_pdf(OUT_DIR / "single_column.pdf")
    build_two_column_pdf(OUT_DIR / "two_column.pdf")
    build_varied_headers_pdf(OUT_DIR / "varied_headers.pdf")
    build_resume_docx(OUT_DIR / "resume.docx")
    build_corrupt_files(OUT_DIR / "corrupt.pdf", OUT_DIR / "corrupt.docx")
    print(f"Fixtures written to {OUT_DIR}")
