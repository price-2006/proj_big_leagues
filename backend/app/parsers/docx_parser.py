"""DOCX -> raw text + layout metadata, via python-docx (see docs/ARCHITECTURE.md §6).

DOCX has no page-position concept the way PDF does; its layout signal is
structural instead — paragraph style name ("Heading 1", "Normal", ...) plus
run-level bold/font-size, which is what the Phase 2 section detector reads
in place of PDF's bbox/font-size heuristics.

python-docx (and the lxml it's built on) is imported lazily, inside
parse_docx() itself, rather than at module level: it's a native-extension
dependency that other, unrelated code paths (JD parsing, matching,
skills browsing) shouldn't have to pay the import cost of, or fail to
boot over, just because something imported this module.

XXE defense (Phase 14, docs/ARCHITECTURE.md §12): python-docx has no
parser-hardening hook of its own — it hands its zip's XML parts straight
to lxml with entity resolution on. `_reject_xxe_payloads` pre-scans every
XML/rels part in the zip — via the stdlib `zipfile` (defusedxml's own
`defusedxml.zipfile.DefusedZipFile` doesn't exist in the installed 0.7.1
release, confirmed by actually importing it, not assumed from docs) —
parsing each part's bytes with `defusedxml.ElementTree` (which raises on
a DOCTYPE or external-entity declaration) *before* python-docx ever
touches them, so a malicious part is rejected up front rather than
relying on python-docx's own (nonexistent) hardening. Each part's
declared (pre-decompression) size is also capped, as a cheap guard
against a decompression-bomb-style member — stdlib `zipfile` doesn't
enforce one itself.
"""
from __future__ import annotations

import zipfile
from io import BytesIO
from typing import TYPE_CHECKING

from defusedxml.common import DefusedXmlException

from app.parsers.exceptions import DocumentParseError
from app.schemas.document import ParsedDocument, TextLine

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph

_MAX_DECOMPRESSED_PART_BYTES = 20 * 1024 * 1024  # generous for a real docx XML part


def _reject_xxe_payloads(data: bytes) -> None:
    from defusedxml import ElementTree as DefusedElementTree

    try:
        with zipfile.ZipFile(BytesIO(data)) as zf:
            for info in zf.infolist():
                if not info.filename.endswith((".xml", ".rels")):
                    continue
                if info.file_size > _MAX_DECOMPRESSED_PART_BYTES:
                    raise DocumentParseError(f"DOCX part '{info.filename}' exceeds the safe decompressed size limit")
                with zf.open(info) as part:
                    DefusedElementTree.parse(part)
    except DefusedXmlException as exc:
        raise DocumentParseError(f"Rejected DOCX containing disallowed XML content: {exc}") from exc
    except zipfile.BadZipFile as exc:
        raise DocumentParseError(f"Could not open DOCX: {exc}") from exc
    except DocumentParseError:
        raise
    except Exception as exc:
        # Anything else (an encrypted zip entry, a malformed-but-not-quite-
        # BadZipFile archive, ...) fails closed the same way the rest of
        # this module does: a clean DocumentParseError, never a raw
        # traceback reaching the API layer.
        raise DocumentParseError(f"Could not safely inspect DOCX contents: {exc}") from exc


def parse_docx(data: bytes) -> ParsedDocument:
    import docx

    _reject_xxe_payloads(data)

    try:
        document = docx.Document(BytesIO(data))
    except Exception as exc:
        raise DocumentParseError(f"Could not open DOCX: {exc}") from exc

    lines: list[TextLine] = []
    text_parts: list[str] = []
    try:
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            font_size, bold = _run_style(paragraph)
            lines.append(
                TextLine(
                    text=text,
                    font_size=font_size,
                    bold=bold,
                    style_name=paragraph.style.name if paragraph.style else None,
                )
            )
            text_parts.append(text)
    except Exception as exc:
        raise DocumentParseError(f"Failed while extracting DOCX text: {exc}") from exc

    if not text_parts:
        raise DocumentParseError("No extractable text found in DOCX")

    return ParsedDocument(file_type="docx", raw_text="\n".join(text_parts), lines=lines)


def _run_style(paragraph: Paragraph) -> tuple[float | None, bool]:
    """Font size/bold come from the runs; fall back to the paragraph style
    when no run overrides it (most resume body text doesn't)."""
    sizes = [run.font.size.pt for run in paragraph.runs if run.font.size is not None]
    bold = any(run.bold for run in paragraph.runs if run.bold is not None)

    style_font = paragraph.style.font if paragraph.style else None
    font_size = max(sizes) if sizes else (style_font.size.pt if style_font and style_font.size else None)
    if not bold and style_font is not None:
        bold = bool(style_font.bold)

    return font_size, bold
